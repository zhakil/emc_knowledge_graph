"""
EMC文件处理服务
支持多种格式的EMC相关文档解析和实体提取
"""

import asyncio
import json
import logging
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple, Union, AsyncGenerator
from dataclasses import dataclass, asdict
from datetime import datetime
import hashlib

# 文件处理依赖
import pandas as pd
import pdfplumber
import docx
from openpyxl import load_workbook
import xml.etree.ElementTree as ET
import aiofiles
import chardet

# 自定义模块
from .content_extractor import EMCContentExtractor
from .format_converter import FormatConverter
from ..ai_integration.deepseek_service import DeepSeekEMCService
from ..knowledge_graph.graph_manager import EMCGraphManager


@dataclass
class FileMetadata:
    """文件元数据"""
    file_id: str
    filename: str
    file_type: str
    size_bytes: int
    mime_type: str
    encoding: Optional[str]
    checksum: str
    upload_time: datetime
    processed: bool = False
    extraction_status: str = "pending"  # pending, processing, completed, failed


@dataclass
class ExtractionResult:
    """提取结果"""
    file_id: str
    entities: List[Dict[str, Any]]
    relationships: List[Dict[str, Any]]
    content_summary: str
    confidence_score: float
    processing_time: float
    extracted_at: datetime


class EMCFileProcessor:
    """EMC文件处理核心类"""
    
    # 支持的文件格式
    SUPPORTED_FORMATS = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.csv': 'text/csv',
        '.json': 'application/json',
        '.xml': 'application/xml',
        '.txt': 'text/plain'
    }
    
    def __init__(
        self, 
        deepseek_service: DeepSeekEMCService, # Assuming DeepSeekEMCService is correctly typed
        storage_path: str = "./uploads",
        graph_manager: Optional[EMCGraphManager] = None # New argument
    ):
        self.deepseek = deepseek_service
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(exist_ok=True)
        self.logger = logging.getLogger(__name__)
        
        if graph_manager:
            self.graph_manager = graph_manager
            self.logger.info("EMCFileProcessor initialized with provided EMCGraphManager.")
        else:
            try:
                self.graph_manager = EMCGraphManager() # EMCGraphManager might need its own dependencies
                self.logger.info("EMCFileProcessor automatically initialized a new EMCGraphManager.")
            except Exception as e:
                self.logger.error(f"Failed to auto-initialize EMCGraphManager in EMCFileProcessor: {e}", exc_info=True)
                self.graph_manager = None

        self.content_extractor = EMCContentExtractor()
        self.format_converter = FormatConverter()
        
        self._processing_stats = {
            'files_processed': 0,
            'entities_extracted': 0,
            'relationships_found': 0,
            'processing_errors': 0,
            'graph_processing_invocation_errors': 0, # Renamed for clarity
            'graph_processing_content_errors': 0
        }
    
    async def process_file(
        self, 
        file_path: Union[str, Path],
        file_id: Optional[str] = None,
        trigger_graph_processing: bool = True # New flag
    ) -> Tuple[FileMetadata, Optional[ExtractionResult]]:
        """
        处理单个文件
        返回元数据和提取结果
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 生成文件ID
        if not file_id:
            file_id = self._generate_file_id(file_path)
        
        # 提取文件元数据
        metadata = await self._extract_metadata(file_path, file_id)
        
        # 验证文件格式
        if not self._is_supported_format(file_path):
            metadata.extraction_status = "unsupported_format"
            return metadata, None
        
        extraction_result: Optional[ExtractionResult] = None

        try:
            # 提取文件内容
            content = await self._extract_content(file_path, metadata.mime_type)
            
            if not content.strip():
                metadata.extraction_status = "empty_content"
                return metadata, None
            
            # --- Current AI extraction (for ExtractionResult) ---
            # This part can remain for now, but its output (ExtractionResult) might be
            # less critical if graph_manager handles the primary structured data output.
            try:
                extraction_result = await self._extract_entities_with_ai(
                    file_id, content, metadata # metadata is FileMetadata object
                )
                # Statistics for this extraction are handled within _extract_entities_with_ai
            except Exception as e:
                self.logger.error(f"Original _extract_entities_with_ai failed for {file_id}: {e}", exc_info=True)
                # self._processing_stats['ai_invocation_errors'] += 1 # Assuming this was the old stat name
                if extraction_result is None:
                    extraction_result = ExtractionResult(
                        file_id=file_id, entities=[], relationships=[], content_summary="AI extraction failed.",
                        confidence_score=0.0, processing_time=0.0, extracted_at=datetime.now()
                    )

            # --- New: EMCGraphManager processing ---
            if trigger_graph_processing:
                if self.graph_manager:
                    self.logger.info(f"Calling EMCGraphManager to process document: {file_id}")
                    try:
                        # Convert FileMetadata object to dict for graph_manager
                        metadata_dict = asdict(metadata) # Ensure asdict is imported from dataclasses
                        graph_summary = await self.graph_manager.process_document_content(
                            text_content=content,
                            document_id=file_id,
                            document_metadata=metadata_dict
                        )
                        self.logger.info(f"Graph processing summary for {file_id}: {graph_summary.get('status')}")
                        if graph_summary.get("status") == "failed" or graph_summary.get("errors"):
                            self._processing_stats['graph_processing_content_errors'] += 1
                            # Update status only if it's not already a more severe failure
                            if not metadata.extraction_status or "failed" not in metadata.extraction_status.lower():
                                metadata.extraction_status = "completed_with_graph_errors"
                        # Optionally, you could merge entities/rels from graph_summary into extraction_result if needed

                    except Exception as e:
                        self.logger.error(f"EMCGraphManager invocation failed for {file_id}: {e}", exc_info=True)
                        self._processing_stats['graph_processing_invocation_errors'] += 1
                        if not metadata.extraction_status or "failed" not in metadata.extraction_status.lower():
                             metadata.extraction_status = "graph_processing_failed"
                else:
                    self.logger.warning(f"Graph manager not available for document {file_id}. Skipping graph processing.")
                    if not metadata.extraction_status or "failed" not in metadata.extraction_status.lower():
                        metadata.extraction_status = "graph_processing_skipped_no_manager"


            metadata.processed = True
            self._processing_stats['files_processed'] += 1
            # Final status update if not set by graph processing stages
            if not metadata.extraction_status or metadata.extraction_status == "pending":
                 if extraction_result and extraction_result.entities: # From old method
                      metadata.extraction_status = "completed_basic_extraction_only"
                 else: # This implies extraction_result might be None or have no entities
                      metadata.extraction_status = "completed_no_entities_found"
            
            return metadata, extraction_result
            
        except Exception as e:
            self.logger.error(f"文件处理失败: {file_path}, 错误: {str(e)}", exc_info=True) # Original error logging
            metadata.extraction_status = "failed_in_file_processing" # More specific outer failure
            self._processing_stats['processing_errors'] += 1
            return metadata, None
    
    async def batch_process_files(
        self, 
        file_paths: List[Union[str, Path]],
        max_concurrent: int = 3
    ) -> AsyncGenerator[Tuple[FileMetadata, Optional[ExtractionResult]], None]:
        """
        批量处理文件
        使用异步并发处理提高效率
        """
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def process_single_file(file_path):
            async with semaphore:
                return await self.process_file(file_path)
        
        # 创建任务列表
        tasks = [process_single_file(fp) for fp in file_paths]
        
        # 并发执行并逐个返回结果
        for task in asyncio.as_completed(tasks):
            yield await task
    
    def _generate_file_id(self, file_path: Path) -> str:
        """生成唯一文件ID"""
        content_hash = hashlib.md5(f"{file_path.name}_{file_path.stat().st_mtime}".encode()).hexdigest()
        return f"file_{content_hash[:12]}"
    
    async def _extract_metadata(self, file_path: Path, file_id: str) -> FileMetadata:
        """提取文件元数据"""
        stat = file_path.stat()
        
        # 检测MIME类型
        mime_type, _ = mimetypes.guess_type(str(file_path))
        if not mime_type:
            suffix = file_path.suffix.lower()
            mime_type = self.SUPPORTED_FORMATS.get(suffix, 'application/octet-stream')
        
        # 计算文件校验和
        checksum = await self._calculate_checksum(file_path)
        
        # 检测文件编码（对文本文件）
        encoding = None
        if mime_type.startswith('text/'):
            encoding = await self._detect_encoding(file_path)
        
        return FileMetadata(
            file_id=file_id,
            filename=file_path.name,
            file_type=file_path.suffix[1:].lower(),
            size_bytes=stat.st_size,
            mime_type=mime_type,
            encoding=encoding,
            checksum=checksum,
            upload_time=datetime.fromtimestamp(stat.st_mtime)
        )
    
    async def _calculate_checksum(self, file_path: Path) -> str:
        """计算文件SHA256校验和"""
        sha256_hash = hashlib.sha256()
        
        async with aiofiles.open(file_path, 'rb') as f:
            while chunk := await f.read(8192):
                sha256_hash.update(chunk)
        
        return sha256_hash.hexdigest()
    
    async def _detect_encoding(self, file_path: Path) -> Optional[str]:
        """检测文件编码"""
        try:
            async with aiofiles.open(file_path, 'rb') as f:
                raw_data = await f.read(10240)  # 读取前10KB检测编码
                
            result = chardet.detect(raw_data)
            return result.get('encoding')
        except Exception:
            return None
    
    def _is_supported_format(self, file_path: Path) -> bool:
        """检查文件格式是否支持"""
        return file_path.suffix.lower() in self.SUPPORTED_FORMATS
    
    async def _extract_content(self, file_path: Path, mime_type: str) -> str:
        """根据文件类型提取内容"""
        try:
            if mime_type == 'application/pdf':
                return await self._extract_pdf_content(file_path)
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return await self._extract_docx_content(file_path)
            elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                return await self._extract_xlsx_content(file_path)
            elif mime_type == 'text/csv':
                return await self._extract_csv_content(file_path)
            elif mime_type == 'application/json':
                return await self._extract_json_content(file_path)
            elif mime_type == 'application/xml':
                return await self._extract_xml_content(file_path)
            elif mime_type == 'text/plain':
                return await self._extract_text_content(file_path)
            else:
                raise ValueError(f"不支持的文件类型: {mime_type}")
                
        except Exception as e:
            self.logger.error(f"内容提取失败: {file_path}, 错误: {str(e)}")
            raise
    
    async def _extract_pdf_content(self, file_path: Path) -> str:
        """提取PDF内容"""
        content_parts = []
        
        def extract_pdf():
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content_parts.append(text)
                    
                    # 提取表格
                    tables = page.extract_tables()
                    for table in tables:
                        table_text = "\n".join(["\t".join(row) for row in table if row])
                        if table_text:
                            content_parts.append(f"\n[表格数据]\n{table_text}\n")
        
        # 在线程池中运行CPU密集的PDF处理
        await asyncio.get_event_loop().run_in_executor(None, extract_pdf)
        
        return "\n\n".join(content_parts)
    
    async def _extract_docx_content(self, file_path: Path) -> str:
        """提取Word文档内容"""
        def extract_docx():
            doc = docx.Document(file_path)
            content_parts = []
            
            # 提取段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
            
            # 提取表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append("\t".join(row_data))
                
                if table_data:
                    content_parts.append(f"\n[表格数据]\n" + "\n".join(table_data) + "\n")
            
            return "\n\n".join(content_parts)
        
        return await asyncio.get_event_loop().run_in_executor(None, extract_docx)
    
    async def _extract_xlsx_content(self, file_path: Path) -> str:
        """提取Excel内容"""
        def extract_xlsx():
            workbook = load_workbook(file_path, read_only=True)
            content_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content_parts.append(f"\n[工作表: {sheet_name}]\n")
                
                # 转换为DataFrame进行处理
                data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        data.append([str(cell) if cell is not None else "" for cell in row])
                
                if data:
                    df = pd.DataFrame(data[1:], columns=data[0] if data else [])
                    content_parts.append(df.to_string(index=False))
            
            workbook.close()
            return "\n\n".join(content_parts)
        
        return await asyncio.get_event_loop().run_in_executor(None, extract_xlsx)
    
    async def _extract_csv_content(self, file_path: Path) -> str:
        """提取CSV内容"""
        def extract_csv():
            # 尝试不同编码
            for encoding in ['utf-8', 'gbk', 'latin1']:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    return f"CSV数据 ({df.shape[0]}行 x {df.shape[1]}列):\n\n{df.to_string(index=False)}"
                except UnicodeDecodeError:
                    continue
            
            # 如果所有编码都失败，尝试自动检测
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding = chardet.detect(raw_data)['encoding']
                
            df = pd.read_csv(file_path, encoding=encoding)
            return f"CSV数据 ({df.shape[0]}行 x {df.shape[1]}列):\n\n{df.to_string(index=False)}"
        
        return await asyncio.get_event_loop().run_in_executor(None, extract_csv)
    
    async def _extract_json_content(self, file_path: Path) -> str:
        """提取JSON内容"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
            
        try:
            data = json.loads(content)
            # 格式化JSON输出
            return f"JSON数据结构:\n\n{json.dumps(data, ensure_ascii=False, indent=2)}"
        except json.JSONDecodeError as e:
            return f"JSON格式错误: {str(e)}\n原始内容:\n{content}"
    
    async def _extract_xml_content(self, file_path: Path) -> str:
        """提取XML内容"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
        
        try:
            root = ET.fromstring(content)
            
            def xml_to_text(element, depth=0):
                result = []
                indent = "  " * depth
                
                if element.text and element.text.strip():
                    result.append(f"{indent}{element.tag}: {element.text.strip()}")
                else:
                    result.append(f"{indent}{element.tag}:")
                
                for child in element:
                    result.extend(xml_to_text(child, depth + 1))
                
                return result
            
            text_content = "\n".join(xml_to_text(root))
            return f"XML结构数据:\n\n{text_content}"
            
        except ET.ParseError as e:
            return f"XML解析错误: {str(e)}\n原始内容:\n{content}"
    
    async def _extract_text_content(self, file_path: Path) -> str:
        """提取纯文本内容"""
        encoding = await self._detect_encoding(file_path) or 'utf-8'
        
        async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
            return await f.read()
    
    async def _extract_entities_with_ai(
        self, 
        file_id: str, 
        content: str, 
        metadata: FileMetadata
    ) -> ExtractionResult:
        """使用AI提取实体和关系"""
        start_time = datetime.now()
        
        try:
            # 调用DeepSeek进行实体提取
            response = await self.deepseek.extract_entities_from_text(
                text_content=content,
                session_id=f"file_processing_{file_id}"
            )
            
            # 解析AI响应
            ai_content = response.get('content', '')
            
            # 尝试解析JSON格式的响应
            try:
                extraction_data = json.loads(ai_content)
            except json.JSONDecodeError:
                # 如果不是JSON格式，使用简单解析
                extraction_data = self._fallback_entity_extraction(content)
            
            # 计算置信度分数
            confidence_score = self._calculate_confidence_score(
                extraction_data, content, response
            )
            
            # 生成内容摘要
            content_summary = self._generate_content_summary(content)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            # 更新统计
            self._processing_stats['entities_extracted'] += len(extraction_data.get('entities', []))
            self._processing_stats['relationships_found'] += len(extraction_data.get('relationships', []))
            
            return ExtractionResult(
                file_id=file_id,
                entities=extraction_data.get('entities', []),
                relationships=extraction_data.get('relationships', []),
                content_summary=content_summary,
                confidence_score=confidence_score,
                processing_time=processing_time,
                extracted_at=datetime.now()
            )
            
        except Exception as e:
            self.logger.error(f"AI实体提取失败: {str(e)}")
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return ExtractionResult(
                file_id=file_id,
                entities=[],
                relationships=[],
                content_summary="提取失败",
                confidence_score=0.0,
                processing_time=processing_time,
                extracted_at=datetime.now()
            )
    
    def _fallback_entity_extraction(self, content: str) -> Dict[str, List]:
        """备用实体提取方法（基于规则）"""
        # 简单的EMC实体识别
        entities = []
        relationships = []
        
        # EMC标准识别
        import re
        standard_patterns = [
            r'EN\s+\d+[-:]\d+',
            r'IEC\s+\d+[-:]\d+',
            r'FCC\s+Part\s+\d+',
            r'CISPR\s+\d+',
            r'MIL-STD-\d+'
        ]
        
        for pattern in standard_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            for match in matches:
                entities.append({
                    'type': 'EMCStandard',
                    'name': match,
                    'properties': {'source': 'rule_based_extraction'}
                })
        
        # 频率范围识别
        freq_pattern = r'(\d+(?:\.\d+)?)\s*([kMG]?Hz)\s*[-到至]\s*(\d+(?:\.\d+)?)\s*([kMG]?Hz)'
        freq_matches = re.findall(freq_pattern, content, re.IGNORECASE)
        for match in freq_matches:
            freq_range = f"{match[0]}{match[1]}-{match[2]}{match[3]}"
            entities.append({
                'type': 'FrequencyRange',
                'name': freq_range,
                'properties': {'source': 'rule_based_extraction'}
            })
        
        return {
            'entities': entities,
            'relationships': relationships
        }
    
    def _calculate_confidence_score(
        self, 
        extraction_data: Dict, 
        content: str, 
        ai_response: Dict
    ) -> float:
        """计算提取结果的置信度分数"""
        score = 0.0
        
        # 基于实体数量
        entity_count = len(extraction_data.get('entities', []))
        if entity_count > 0:
            score += min(entity_count * 0.1, 0.4)
        
        # 基于关系数量
        rel_count = len(extraction_data.get('relationships', []))
        if rel_count > 0:
            score += min(rel_count * 0.15, 0.3)
        
        # 基于AI响应质量
        if ai_response.get('usage', {}).get('total_tokens', 0) > 100:
            score += 0.2
        
        # 基于内容长度
        if len(content) > 1000:
            score += 0.1
        
        return min(score, 1.0)
    
    def _generate_content_summary(self, content: str, max_length: int = 200) -> str:
        """生成内容摘要"""
        # 简单的摘要生成
        sentences = content.split('.')[:3]  # 取前3个句子
        summary = '. '.join(sentences).strip()
        
        if len(summary) > max_length:
            summary = summary[:max_length-3] + '...'
        
        return summary
    
    def get_processing_stats(self) -> Dict[str, int]:
        """获取处理统计信息"""
        return self._processing_stats.copy()
    
    async def cleanup_temp_files(self, max_age_days: int = 7):
        """清理临时文件"""
        import time
        current_time = time.time()
        max_age_seconds = max_age_days * 24 * 3600
        
        temp_dir = self.storage_path / "temp"
        if temp_dir.exists():
            for file_path in temp_dir.iterdir():
                if file_path.is_file():
                    file_age = current_time - file_path.stat().st_mtime
                    if file_age > max_age_seconds:
                        file_path.unlink()
                        self.logger.info(f"清理临时文件: {file_path}")


class EMCContentExtractor:
    """EMC内容提取器"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def extract_emc_keywords(self, content: str) -> List[str]:
        """提取EMC关键词"""
        # EMC专业术语词典
        emc_keywords = [
            'electromagnetic compatibility', 'emc', 'emi', 'conducted emission',
            'radiated emission', 'immunity', 'susceptibility', 'surge',
            'burst', 'flicker', 'harmonics', 'electrostatic discharge',
            'esd', 'voltage dip', 'power interruption'
        ]
        
        found_keywords = []
        content_lower = content.lower()
        
        for keyword in emc_keywords:
            if keyword in content_lower:
                found_keywords.append(keyword)
        
        return found_keywords


class FormatConverter:
    """格式转换器"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    async def convert_to_standard_format(
        self, 
        content: str, 
        source_format: str, 
        target_format: str = 'json'
    ) -> str:
        """转换为标准格式"""
        if target_format == 'json':
            return json.dumps({
                'content': content,
                'source_format': source_format,
                'converted_at': datetime.now().isoformat()
            }, ensure_ascii=False, indent=2)
        
        return content


# 工厂函数
def create_emc_file_processor(
    deepseek_service: DeepSeekEMCService,
    storage_path: str = "./uploads",
    graph_manager: Optional[EMCGraphManager] = None # Add graph_manager
) -> EMCFileProcessor:
    """创建EMC文件处理器实例"""
    return EMCFileProcessor(
        deepseek_service,
        storage_path,
        graph_manager=graph_manager # Pass it to constructor
    )